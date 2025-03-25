import streamlit as st
import requests
from docx import Document

UNIT_INSIGHTS_EVALUATOR = st.secrets["UNIT_INSIGHTS_EVALUATOR"]
SUMMARIZED_REPORT_EVALUATOR = st.secrets["SUMMARIZED_REPORT_EVALUATOR"]

# Khởi tạo state nếu chưa tồn tại
if 'summarized_report_content' not in st.session_state:
    st.session_state.summarized_report_content = None
if 'outline_content' not in st.session_state:
  st.session_state.outline_content = None
if 'unit_report' not in st.session_state:
  st.session_state.unit_report = None
if 'unit_report_insights' not in st.session_state:
  st.session_state.unit_report_insights = None

tab1, tab2 = st.tabs(["Unit Insights Evaluation", "Summarized Report Evaluation"])

with tab1:
    st.title("Tải lên báo cáo đơn vị và insights báo cáo")
    col1, col2 = st.columns(2)
    with col1:
        uploaded_unit_report = st.file_uploader("Tải lên Báo cáo đơn vị", type=["docx", "txt"])
        if uploaded_unit_report is not None:
            if uploaded_unit_report.type == "text/plain":
                st.session_state.unit_report_content = uploaded_unit_report.getvalue().decode("utf-8")
            else:
                report = Document(uploaded_unit_report)
                st.session_state.unit_report_content = '\n'.join([para.text for para in report.paragraphs])
        if uploaded_unit_report is not None:
            st.write("Báo cáo đã được tải lên thành công!")

    with col2:
        uploaded_unit_report_insights = st.file_uploader("Tải lên Insights báo cáo", type=["docx", "txt"])
        if uploaded_unit_report_insights is not None:
            if uploaded_unit_report_insights.type == "text/plain":
                st.session_state.unit_report_insights_content = uploaded_unit_report_insights.getvalue().decode("utf-8")
            else:
                report = Document(uploaded_unit_report_insights)
                st.session_state.unit_report_insights_content = '\n'.join([para.text for para in report.paragraphs])
        if uploaded_unit_report_insights is not None:
            st.write("Insights báo cáo đã được tải lên thành công!")

    if st.button("Evaluate Unit Report Insights"):
        if uploaded_unit_report is not None:
            payload = {
                "question": "So sánh các insights báo cáo đối với báo cáo đơn vị",
                "overrideConfig": {
                    "promptValues": {
                        "unit_report": st.session_state.unit_report_content,
                        "unit_report_insights": st.session_state.unit_report_insights_content
                    }
                }
            }
            with st.spinner("Đang đánh giá report insights..."):
                response = requests.post(UNIT_INSIGHTS_EVALUATOR, json=payload)
            try:
                st.markdown(response.json()["text"])
            except:
                st.write(response.json())

with tab2:
    st.title("Tải lên báo cáo tổng hợp và đề cương báo cáo")
    col3, col4 = st.columns(2)
    with col3:
        uploaded_summarized_report = st.file_uploader("Tải lên báo cáo tổng hợp", type=["docx", "txt"])
        if uploaded_summarized_report is not None:
            if uploaded_summarized_report.type == "text/plain":
                st.session_state.summarized_report_content = uploaded_summarized_report.getvalue().decode("utf-8")
            else:
                report = Document(uploaded_summarized_report)
                st.session_state.summarized_report_content = '\n'.join([para.text for para in report.paragraphs])
        if uploaded_summarized_report is not None:
            st.write("Báo cáo đã được tải lên thành công!")

    with col4:
        uploaded_outline = st.file_uploader("Tải lên Đề cương báo cáo", type=["docx", "txt"])
        if uploaded_outline is not None:
            if uploaded_outline.type == "text/plain":
                st.session_state.outline_content = uploaded_outline.getvalue().decode("utf-8")
            else:
                outline = Document(uploaded_outline)
                st.session_state.outline_content = '\n'.join([para.text for para in outline.paragraphs])
        if uploaded_outline is not None:
            st.write("Đề cương báo cáo đã được tải lên thành công!")

    #  Tạo nút để tải xuống file
    if st.button("Evaluate Summarized Report"):
        if uploaded_summarized_report is not None:
            payload = {
                "question": "So sánh nội dung báo cáo với đề cương báo cáo",
                "overrideConfig": {
                    "promptValues": {
                        "report": st.session_state.summarized_report_content,
                        "report_template": st.session_state.outline_content
                    }
                }
            }
            with st.spinner("Đang đánh giá báo cáo..."):
                response = requests.post(SUMMARIZED_REPORT_EVALUATOR, json=payload)
            try:
                st.markdown(response.json()["text"])
            except:
                st.write(response.json())
    
if st.button("Reset"):
    st.session_state.summarized_report_content = None
    st.session_state.outline_content = None
    st.session_state.unit_report_content = None
    st.session_state.unit_report_insights_content = None
    st.rerun()